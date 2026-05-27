# -*- coding: utf-8 -*-
from docx import Document
from docx.oxml import OxmlElement
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.text.paragraph import Paragraph as DocxParagraph

doc = Document('Manual_Validacao_CISPR15_LABELO_converted.docx')

# ── helpers ───────────────────────────────────────────────────────────────────
def ins_after(anchor):
    new_p = OxmlElement('w:p')
    anchor._p.addnext(new_p)
    return DocxParagraph(new_p, anchor._p.getparent())

def new_bullet(anchor, text, size=10.5):
    p = ins_after(anchor)
    p.paragraph_format.left_indent  = Cm(1.0)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    run = p.add_run('•  ' + text)
    run.font.name = 'Arial'
    run.font.size = Pt(size)
    return p

def new_numbered(anchor, num, text, size=10.5):
    p = ins_after(anchor)
    p.paragraph_format.left_indent  = Cm(1.0)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    run = p.add_run(str(num) + '.  ' + text)
    run.font.name = 'Arial'
    run.font.size = Pt(size)
    return p

# ── localiza âncoras ─────────────────────────────────────────────────────────
DARK  = RGBColor(0x1F, 0x38, 0x64)

anchor_32 = None
anchor_33 = None
for p in doc.paragraphs:
    if 'cispr15_clientes.json e cispr15_relatorios.json' in p.text and anchor_32 is None:
        anchor_32 = p
    if 'Pasta dedicada para o arquivo cispr15_agenda.json' in p.text and anchor_33 is None:
        anchor_33 = p

assert anchor_32 and anchor_33, 'Paragrafos nao encontrados'

# ═════════════════════════════════════════════════════════════════════════════
# INSERIR APÓS 3.2 — pasta local padrão + migração para rede
# ═════════════════════════════════════════════════════════════════════════════

# Título "Pasta local padrão"
p = ins_after(anchor_32)
p.paragraph_format.space_before = Pt(6)
p.paragraph_format.space_after  = Pt(2)
run = p.add_run('Pasta local padrão (campo vazio):')
run.bold = True
run.font.name = 'Arial'
run.font.size = Pt(10.5)
run.font.color.rgb = DARK

# Explicação
p = ins_after(p)
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p.paragraph_format.space_before = Pt(1)
p.paragraph_format.space_after  = Pt(2)
run = p.add_run(
    'Quando o campo "Pasta de Dados Compartilhados" estiver vazio, o sistema utiliza '
    'automaticamente a seguinte pasta no computador local, independente de onde o '
    'aplicativo está instalado:'
)
run.font.name = 'Arial'
run.font.size = Pt(10.5)

# Caminho em destaque
caminho = 'Documentos\\CISPR 15 LABELO\\dados\\'
p = ins_after(p)
p.paragraph_format.left_indent  = Cm(1.5)
p.paragraph_format.space_before = Pt(3)
p.paragraph_format.space_after  = Pt(3)
run = p.add_run(caminho)
run.bold = True
run.font.name = 'Courier New'
run.font.size = Pt(10)

# Arquivos
p = ins_after(p)
p.paragraph_format.space_before = Pt(2)
p.paragraph_format.space_after  = Pt(1)
run = p.add_run('Os três arquivos de dados ficam nessa pasta:')
run.font.name = 'Arial'
run.font.size = Pt(10.5)

b1 = new_bullet(p,  'cispr15_clientes.json — cadastro de clientes do laboratório')
b2 = new_bullet(b1, 'cispr15_relatorios.json — metadados dos relatórios emitidos')
b3 = new_bullet(b2, 'cispr15_agenda.json — itens da agenda de execução')

# Título migração
p = ins_after(b3)
p.paragraph_format.space_before = Pt(8)
p.paragraph_format.space_after  = Pt(2)
run = p.add_run('Como migrar para pasta de rede (compartilhamento entre computadores):')
run.bold = True
run.font.name = 'Arial'
run.font.size = Pt(10.5)
run.font.color.rgb = DARK

rede_exemplo = r'\\servidor\projetos\CISPR15\dados'
n1 = new_numbered(p, 1,
    'Copie os três arquivos de Documentos\\CISPR 15 LABELO\\dados\\ para a pasta de rede '
    'desejada (ex.: ' + rede_exemplo + ').')
n2 = new_numbered(n1, 2,
    'Abra Configurações no aplicativo e informe o caminho de rede no campo '
    '"Pasta de Dados Compartilhados".')
n3 = new_numbered(n2, 3,
    'Clique em Salvar. A partir deste momento todos os computadores configurados com o '
    'mesmo caminho compartilham o mesmo banco de dados.')

# Nota backup
p = ins_after(n3)
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p.paragraph_format.space_before = Pt(4)
p.paragraph_format.space_after  = Pt(6)
run = p.add_run(
    'Observação: a pasta local padrão não é apagada após a migração e pode ser usada '
    'como backup em caso de perda de acesso ao servidor de rede.'
)
run.italic = True
run.font.name = 'Arial'
run.font.size = Pt(10)

# ═════════════════════════════════════════════════════════════════════════════
# INSERIR APÓS 3.3 — nota sobre caminho padrão da agenda
# ═════════════════════════════════════════════════════════════════════════════

p33 = ins_after(anchor_33)
p33.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p33.paragraph_format.space_before = Pt(2)
p33.paragraph_format.space_after  = Pt(6)
run = p33.add_run(
    'Quando deixada vazia, utiliza a mesma pasta configurada em 3.2 acima — ou, se ambas '
    'estiverem vazias, a pasta padrão local Documentos\\CISPR 15 LABELO\\dados\\. '
    'Pode ser configurada separadamente para que técnicos que apenas consultam a agenda '
    'tenham acesso somente ao arquivo cispr15_agenda.json, sem exposição do banco de '
    'clientes e relatórios — útil em ambientes onde diferentes perfis têm permissões '
    'distintas na rede.'
)
run.font.name = 'Arial'
run.font.size = Pt(10.5)

# ── salva ─────────────────────────────────────────────────────────────────────
doc.save('Manual_Validacao_CISPR15_LABELO_converted.docx')
print('Salvo com sucesso.')
