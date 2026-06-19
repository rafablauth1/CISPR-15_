# -*- coding: utf-8 -*-
"""
Gera o documento de Validação / Manual do Sistema CISPR 15 LABELO
para o Sistema da Qualidade do LABELO-PUCRS.
"""
import os, datetime
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

BASE   = Path(r"C:\Users\Notla\OneDrive\Área de Trabalho\DIONATA\cispr15-standalone")
SHOTS  = BASE / "screenshots_manual"
LOGO   = BASE / "public" / "formularios" / "emc" / "pucrs-logo.png"
CRL    = BASE / "public" / "formularios" / "emc" / "crl0075.jpg"
OUT    = BASE / "Manual_Validacao_CISPR15_LABELO.docx"

TODAY  = datetime.date.today().strftime("%d/%m/%Y")

# ─── cores ───────────────────────────────────────────────────────────────
DARK      = RGBColor(0x1F, 0x38, 0x64)   # azul escuro PUCRS
BLUE      = RGBColor(0x1F, 0x38, 0x64)   # azul PUCRS
GRAY1     = RGBColor(0x59, 0x59, 0x59)   # cinza escuro
GRAY2     = RGBColor(0x80, 0x80, 0x80)   # cinza médio
WHITE     = RGBColor(0xFF, 0xFF, 0xFF)
BLACK     = RGBColor(0x00, 0x00, 0x00)
FOOTER_BG = RGBColor(0xF2, 0xF2, 0xF2)

# Hex strings usados em set_cell_bg
BG_DARK   = '1F3864'   # azul escuro PUCRS — faixa do cabeçalho
BG_GRAY1  = 'D9D9D9'   # cinza claro — barras de seção / cabeçalho de tabelas
BG_GRAY2  = 'BDD7EE'   # azul claro — 2ª faixa do cabeçalho
BG_ALT    = 'F2F2F2'   # cinza muito claro — linhas alternadas
BG_FOOTER = 'F2F2F2'   # rodapé

# ─── helpers ─────────────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color: str):
    """Define background color de célula via XML."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color.upper())
    tcPr.append(shd)

def set_cell_borders(cell, **kwargs):
    """Define bordas de uma célula."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tbl_brd = OxmlElement('w:tcBorders')
    for side in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        kw = kwargs.get(side, {'val': 'single', 'sz': 4, 'color': '999999'})
        el = OxmlElement(f'w:{ side }')
        for k, v in kw.items():
            el.set(qn(f'w:{k}'), str(v))
        tbl_brd.append(el)
    tcPr.append(tbl_brd)

def para_space(para, before=0, after=0):
    pPr = para._p.get_or_add_pPr()
    sp  = OxmlElement('w:spacing')
    sp.set(qn('w:before'), str(before))
    sp.set(qn('w:after'),  str(after))
    pPr.append(sp)

def add_run(para, text, bold=False, italic=False, size=None, color=None, font='Arial'):
    run = para.add_run(text)
    run.bold   = bold
    run.italic = italic
    if font:   run.font.name = font
    if size:   run.font.size = Pt(size)
    if color:  run.font.color.rgb = color
    return run

def heading_para(doc, text, level=1, size=12, color=None, bg=None, bold=True, space_before=6, space_after=4):
    """Cria parágrafo de título com fundo opcional."""
    para = doc.add_paragraph()
    para_space(para, before=space_before * 20, after=space_after * 20)
    if bg:
        # usar tabela 1x1 para simular fundo
        tbl = doc.add_table(rows=1, cols=1)
        tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
        cell = tbl.cell(0, 0)
        set_cell_bg(cell, bg)
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)
        p.paragraph_format.left_indent  = Cm(0.3)
        run = p.add_run(text)
        run.bold = bold
        run.font.name = 'Arial'
        run.font.size = Pt(size)
        if color: run.font.color.rgb = color
        return tbl
    else:
        add_run(para, text, bold=bold, size=size, color=color)
        return para

def section_bar(doc, text, size=11.5):
    """Barra de seção cinza — como no relatório."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, BG_DARK)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.left_indent  = Cm(0.3)
    run = p.add_run(text)
    run.bold = True
    run.font.name = 'Arial'
    run.font.size = Pt(size)
    run.font.color.rgb = WHITE
    return tbl

def body_para(doc, text, size=11, justify=True, indent=None, bold=False, italic=False, color=None, before=2, after=2):
    para = doc.add_paragraph()
    para_space(para, before=before * 20, after=after * 20)
    if justify:
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if indent:
        para.paragraph_format.left_indent = Cm(indent)
    add_run(para, text, bold=bold, italic=italic, size=size, color=color)
    return para

def bullet(doc, text, size=11):
    para = doc.add_paragraph(style='List Bullet')
    para_space(para, before=1 * 20, after=1 * 20)
    add_run(para, text, size=size)
    return para

def add_screenshot(doc, filename, caption, max_width_cm=15.5):
    """Insere screenshot com legenda."""
    path = SHOTS / filename
    if not path.exists():
        body_para(doc, f'[Imagem não encontrada: {filename}]', italic=True, color=RGBColor(0xAA, 0x00, 0x00))
        return
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para_space(para, before=2 * 20, after=1 * 20)
    run = para.add_run()
    run.add_picture(str(path), width=Cm(max_width_cm))

    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para_space(cap, before=1 * 20, after=6 * 20)
    add_run(cap, f'Figura — {caption}', italic=True, size=9, color=BLACK)

def page_break(doc):
    doc.add_page_break()

# ─── cabeçalho do documento (header Word) ────────────────────────────────
def build_header(doc):
    """
    Cria o cabeçalho Word replicando o design do relatório:
    faixa escura com logo PUCRS + nome lab | sem nº de relatório/amostra.
    """
    section = doc.sections[0]
    section.header_distance = Cm(0.5)
    header  = section.header
    # Limpar parágrafo padrão do header
    for p in header.paragraphs:
        p.clear()

    # Tabela do cabeçalho: 3 colunas [logo | nome | vazio]
    tbl = header.add_table(rows=2, cols=3, width=Cm(18.4))
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    # --- linha 1: fundo escuro ---
    r1c1 = tbl.cell(0, 0)
    r1c2 = tbl.cell(0, 1)
    r1c3 = tbl.cell(0, 2)
    for c in [r1c1, r1c2, r1c3]:
        set_cell_bg(c, BG_DARK)

    # Logo PUCRS
    r1c1.width = Cm(2.2)
    if LOGO.exists():
        p = r1c1.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(3)
        run = p.add_run()
        run.add_picture(str(LOGO), height=Cm(0.9))

    # Nome do laboratório
    r1c2.width = Cm(13.0)
    p2 = r1c2.paragraphs[0]
    p2.paragraph_format.space_before = Pt(3)
    p2.paragraph_format.space_after  = Pt(1)
    p2.paragraph_format.left_indent  = Cm(0.3)
    run2 = p2.add_run('LABELO – Laboratórios Especializados em Eletroeletrônica | Calibração e Ensaios')
    run2.bold = True
    run2.font.name  = 'Arial'
    run2.font.size  = Pt(7.5)
    run2.font.color.rgb = WHITE
    p2b = r1c2.add_paragraph()
    p2b.paragraph_format.left_indent = Cm(0.3)
    p2b.paragraph_format.space_after = Pt(3)
    run2b = p2b.add_run('Pontifícia Universidade Católica do Rio Grande do Sul')
    run2b.font.name  = 'Arial'
    run2b.font.size  = Pt(5.5)
    run2b.font.color.rgb = WHITE

    # Célula direita: tipo do documento
    r1c3.width = Cm(3.2)
    p3 = r1c3.paragraphs[0]
    p3.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p3.paragraph_format.space_before = Pt(4)
    p3.paragraph_format.right_indent = Cm(0.3)
    run3 = p3.add_run('Manual / Validação')
    run3.font.name  = 'Arial'
    run3.font.size  = Pt(6)
    run3.font.color.rgb = WHITE

    # --- linha 2: fundo cinza claro (acreditação + título doc) ---
    r2c1 = tbl.cell(1, 0)
    r2c2 = tbl.cell(1, 1)
    r2c3 = tbl.cell(1, 2)
    merged = r2c1.merge(r2c2).merge(r2c3)
    set_cell_bg(merged, BG_GRAY2)
    pm = merged.paragraphs[0]
    pm.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pm.paragraph_format.space_before = Pt(2)
    pm.paragraph_format.space_after  = Pt(2)
    runm = pm.add_run('Acreditado Cgcre · ABNT NBR ISO/IEC 17025 · CRL 0075    |    Software CISPR 15 LABELO — Manual e Validação do Sistema')
    runm.font.name = 'Arial'
    runm.font.size = Pt(6.5)
    runm.font.color.rgb = BLACK

    # Espaço após tabela
    sp = header.add_paragraph()
    sp.paragraph_format.space_after = Pt(4)

# ─── rodapé do documento ─────────────────────────────────────────────────
def build_footer(doc):
    section = doc.sections[0]
    section.footer_distance = Cm(0.5)
    footer = section.footer
    for p in footer.paragraphs:
        p.clear()

    tbl = footer.add_table(rows=1, cols=3, width=Cm(18.4))
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    c1, c2, c3 = tbl.cell(0, 0), tbl.cell(0, 1), tbl.cell(0, 2)
    for c in [c1, c2, c3]:
        set_cell_bg(c, BG_FOOTER)

    # Borda superior escura via XML (simulada)
    c1.width = Cm(2.5)
    p1 = c1.paragraphs[0]
    p1.paragraph_format.space_before = Pt(3)
    p1.paragraph_format.left_indent  = Cm(0.2)
    run1 = p1.add_run('LABELO\nPUCRS')
    run1.bold = True
    run1.font.name = 'Arial'
    run1.font.size = Pt(7)
    run1.font.color.rgb = BLACK

    c2.width = Cm(13.4)
    p2 = c2.paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_before = Pt(3)
    run2 = p2.add_run(
        'Av. Ipiranga n° 6681, Prédio 30 Bloco A, Sala 210 – Partenon · CEP 90619-900 – Porto Alegre – RS – Brasil\n'
        'Tel.: (51) 3320 3551 · labelo@pucrs.br · www.labelo.com.br'
    )
    run2.font.name = 'Arial'
    run2.font.size = Pt(6.5)
    run2.font.color.rgb = BLACK

    c3.width = Cm(2.5)
    p3 = c3.paragraphs[0]
    p3.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p3.paragraph_format.space_before = Pt(3)
    p3.paragraph_format.right_indent = Cm(0.2)
    # número de página
    fld_run = p3.add_run()
    fld_run.font.name = 'Arial'
    fld_run.font.size = Pt(6.5)
    fld_run.font.color.rgb = BLACK
    # PAGE field
    fldChar1 = OxmlElement('w:fldChar'); fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText'); instrText.text = ' PAGE '; instrText.set(qn('xml:space'), 'preserve')
    fldChar2 = OxmlElement('w:fldChar'); fldChar2.set(qn('w:fldCharType'), 'end')
    for el in [fldChar1, instrText, fldChar2]:
        fld_run._r.append(el)

# ─── página de capa ──────────────────────────────────────────────────────
def build_capa(doc):
    # Logo PUCRS centralizado
    if LOGO.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para_space(p, before=20 * 20, after=6 * 20)
        r = p.add_run()
        r.add_picture(str(LOGO), height=Cm(2.5))

    p0 = doc.add_paragraph()
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para_space(p0, before=4 * 20, after=2 * 20)
    add_run(p0, 'Pontifícia Universidade Católica do Rio Grande do Sul', bold=True, size=13, color=BLUE)

    p1 = doc.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para_space(p1, before=0, after=2 * 20)
    add_run(p1, 'LABELO – Laboratórios Especializados em Eletroeletrônica', bold=True, size=11.5, color=DARK)

    # Linha separadora
    sep = doc.add_paragraph()
    sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para_space(sep, before=8 * 20, after=8 * 20)
    add_run(sep, '─' * 60, size=8, color=GRAY1)

    # Título do documento
    pt = doc.add_paragraph()
    pt.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para_space(pt, before=0, after=4 * 20)
    add_run(pt, 'MANUAL E VALIDAÇÃO DE SOFTWARE', bold=True, size=16, color=BLUE)

    pt2 = doc.add_paragraph()
    pt2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para_space(pt2, before=0, after=4 * 20)
    add_run(pt2, 'Sistema de Geração de Relatórios CISPR 15 LABELO', bold=True, size=13, color=DARK)

    pt3 = doc.add_paragraph()
    pt3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para_space(pt3, before=0, after=30 * 20)
    add_run(pt3, 'Compatibilidade Eletromagnética — Equipamentos de Iluminação Elétrica', size=11, italic=True, color=BLACK)

    # Tabela de controle do documento
    tbl = doc.add_table(rows=7, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = 'Table Grid'
    dados = [
        ('Identificação do Documento', 'MV-EMC-CISPR15-001'),
        ('Revisão',                    '01'),
        ('Data de Emissão',            TODAY),
        ('Elaborado por',              'Dionata Rafael Blauth da Paixão Nunes'),
        ('Revisado por',               'Jonathan Culau'),
        ('Aprovado por',               'Laboratório LABELO-PUCRS'),
        ('Classificação',              'Documento Interno — Sistema da Qualidade'),
    ]
    for i, (label, valor) in enumerate(dados):
        c1 = tbl.cell(i, 0)
        c2 = tbl.cell(i, 1)
        set_cell_bg(c1, BG_GRAY1)
        p1 = c1.paragraphs[0]; p1.paragraph_format.left_indent = Cm(0.2)
        add_run(p1, label, bold=True, size=9.5)
        p2 = c2.paragraphs[0]; p2.paragraph_format.left_indent = Cm(0.2)
        add_run(p2, valor, size=9.5)

    doc.add_page_break()

# ─── corpo do documento ──────────────────────────────────────────────────
def build_body(doc):

    # ══════════════════════════════════════════════════════════════════════
    # 1. INTRODUÇÃO
    # ══════════════════════════════════════════════════════════════════════
    section_bar(doc, '1. Introdução e Objetivo')
    body_para(doc,
        'Este documento descreve o funcionamento e os critérios de validação do software '
        'CISPR 15 LABELO, desenvolvido para uso interno do LABELO-PUCRS (Laboratório de '
        'Ensaios em Compatibilidade Eletromagnética). O sistema é utilizado para a '
        'geração automatizada de relatórios de ensaio conforme a norma NBR IEC/CISPR 15/2014 '
        '— Limites e métodos de medição das radioperturbações de equipamentos elétricos de '
        'iluminação e similares.')
    body_para(doc,
        'O objetivo deste manual é atender aos requisitos do Sistema da Qualidade do laboratório, '
        'documentando as funcionalidades do sistema, os fluxos de trabalho, as configurações de '
        'infraestrutura e os critérios de validação do software, conforme exigido pela '
        'ABNT NBR ISO/IEC 17025:2017.')

    # ══════════════════════════════════════════════════════════════════════
    # 2. INFORMAÇÕES TÉCNICAS
    # ══════════════════════════════════════════════════════════════════════
    section_bar(doc, '2. Informações Técnicas do Sistema')
    body_para(doc, 'A tabela a seguir apresenta as características técnicas do software:')

    tbl = doc.add_table(rows=8, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl.style = 'Table Grid'
    itens = [
        ('Nome do sistema',         'CISPR 15 LABELO'),
        ('Tipo',                    'Aplicação desktop (Electron + Next.js 14 Standalone)'),
        ('Plataforma',              'Microsoft Windows 10/11 x64'),
        ('Linguagem / Framework',   'TypeScript · React 18 · Next.js 14 · Electron 31'),
        ('Geração de PDF',          'Puppeteer-core (Chromium headless via Electron)'),
        ('Leitura de DOCX',         'Mammoth.js (conversão HTML do arquivo Radimation)'),
        ('Planilha de registro',    'XLSX (xlsx.js) — Excel .xlsx em pasta de rede'),
        ('Armazenamento de dados',  r'Arquivos JSON em Documentos\CISPR 15 LABELO\dados\ (padrão) ou pasta de rede configurável'),
    ]
    for i, (k, v) in enumerate(itens):
        set_cell_bg(tbl.cell(i, 0), BG_GRAY1)
        p1 = tbl.cell(i, 0).paragraphs[0]
        p1.paragraph_format.left_indent = Cm(0.2)
        add_run(p1, k, bold=True, size=9.5)
        p2 = tbl.cell(i, 1).paragraphs[0]
        p2.paragraph_format.left_indent = Cm(0.2)
        add_run(p2, v, size=9.5)

    doc.add_paragraph()

    # ══════════════════════════════════════════════════════════════════════
    # 3. CONFIGURAÇÕES — DIRETÓRIOS E SENHA
    # ══════════════════════════════════════════════════════════════════════
    section_bar(doc, '3. Configurações do Sistema — Diretórios e Segurança')
    body_para(doc,
        'O sistema é configurado através da tela "Configurações", acessível pelo ícone de engrenagem '
        'na parte superior direita do formulário principal. As configurações são armazenadas em '
        r'%AppData%\CISPR15-LABELO\settings.json e persistem entre sessões.')

    heading_para(doc, '3.1 Planilha de Registro (Excel)', size=11, bold=True, space_before=8, space_after=3)
    body_para(doc,
        r'Caminho da planilha Excel onde os números de relatório são registrados automaticamente '
        r'ao gerar cada relatório. Pode estar em pasta local ou de rede (ex.: \\servidor\projetos\...). '
        r'O sistema utiliza a biblioteca XLSX para ler e gravar na planilha sem abrir o Excel.')

    heading_para(doc, '3.2 Pasta de Dados Compartilhados', size=11, bold=True, space_before=8, space_after=3)
    body_para(doc,
        r'Pasta onde os arquivos cispr15_clientes.json e cispr15_relatorios.json são armazenados. '
        r'Quando o campo estiver vazio, o sistema utiliza automaticamente a pasta padrão local:')
    body_para(doc,
        'Documentos\\CISPR 15 LABELO\\dados\\',
        bold=True, indent=1)
    body_para(doc,
        r'Este caminho é fixo e idêntico em qualquer computador, independente de onde o aplicativo '
        r'está instalado. O caminho efetivo é sempre exibido como texto de exemplo no próprio campo '
        r'em Configurações. Para compartilhar dados entre computadores do laboratório, siga o '
        r'procedimento abaixo:')
    bullet(doc, r'Copie os três arquivos da pasta padrão (cispr15_agenda.json, cispr15_clientes.json, cispr15_relatorios.json) para a pasta de rede desejada (ex.: \\servidor\projetos\CISPR15\dados).')
    bullet(doc, r'Abra Configurações no aplicativo e informe o caminho de rede no campo "Pasta de Dados Compartilhados".')
    bullet(doc, r'Clique em Salvar. A partir deste momento todos os computadores configurados com o mesmo caminho compartilham o mesmo banco de dados.')
    body_para(doc,
        r'Importante: a pasta local padrão continua existindo como backup histórico após a migração para rede. '
        r'Ela não é apagada e pode ser usada para restauração em caso de perda de acesso ao servidor.',
        italic=True)

    heading_para(doc, '3.3 Pasta da Agenda de Execução', size=11, bold=True, space_before=8, space_after=3)
    body_para(doc,
        r'Pasta dedicada para o arquivo cispr15_agenda.json. Se deixada vazia, utiliza a mesma pasta '
        r'de dados compartilhados configurada acima (ou a pasta padrão local). Pode ser configurada '
        r'separadamente para permitir que técnicos que apenas consultam a agenda tenham acesso somente '
        r'a esta pasta, sem exposição do banco de clientes e relatórios. '
        r'Útil em ambientes onde diferentes perfis de usuário têm permissões distintas na rede.')

    heading_para(doc, '3.4 Pasta de Cópias de PDF', size=11, bold=True, space_before=8, space_after=3)
    body_para(doc,
        'Pasta de destino para cópias automáticas dos PDFs gerados. Toda vez que um relatório '
        'ou emenda é gerado (inclusive após assinatura), uma cópia é salva nesta pasta. '
        'Ao excluir um item da agenda, a cópia correspondente também é removida automaticamente. '
        'Indicada para acesso de consulta sem expor a pasta original da EUT.')

    heading_para(doc, '3.5 Salvamento Automático de PDF na Pasta da EUT', size=11, bold=True, space_before=8, space_after=3)
    body_para(doc,
        'Quando ativada, a opção "Salvar PDF automaticamente na pasta da EUT" salva uma cópia '
        'do relatório gerado diretamente na pasta do ensaio (mesma pasta do arquivo .docx do '
        'Radimation e das fotos), sem necessidade de ação manual pelo operador.')

    heading_para(doc, '3.6 Senha de Acesso à Emissão', size=11, bold=True, space_before=8, space_after=3)
    body_para(doc,
        'O sistema permite configurar uma senha de acesso à área de emissão de relatórios. '
        'Quando configurada, a senha é solicitada a cada nova sessão do aplicativo antes de '
        'permitir o acesso ao formulário de emissão e às configurações. A Agenda de Execução '
        'permanece sempre acessível sem senha, permitindo consulta por qualquer usuário. '
        'A senha é armazenada no arquivo settings.json em %AppData%.')

    add_screenshot(doc, '05_configuracoes.png',
        'Tela de Configurações — Parte superior: planilha, dados compartilhados e pasta da agenda')
    add_screenshot(doc, '05b_configuracoes_inferior.png',
        'Tela de Configurações — Parte inferior: cópias de PDF, salvamento automático e senha de acesso')

    page_break(doc)

    # ══════════════════════════════════════════════════════════════════════
    # 4. AUTENTICAÇÃO — TELA DE SENHA
    # ══════════════════════════════════════════════════════════════════════
    section_bar(doc, '4. Autenticação — Controle de Acesso à Emissão')
    body_para(doc,
        'Quando uma senha de emissão está configurada, ao iniciar uma nova sessão do aplicativo '
        'o usuário é apresentado ao modal de autenticação antes de acessar o formulário de '
        'emissão ou as configurações. O sistema usa sessionStorage para manter a autenticação '
        'durante a sessão; ao fechar e reabrir o aplicativo, a senha é solicitada novamente.')
    body_para(doc,
        'O modal é exibido também ao tentar gerar uma emenda sobre um relatório já bloqueado, '
        'garantindo que apenas pessoal autorizado emita correções em documentos finalizados.')

    add_screenshot(doc, '10_modal_senha.png',
        'Modal de autenticação — solicitação de senha ao iniciar sessão de emissão')

    page_break(doc)

    # ══════════════════════════════════════════════════════════════════════
    # 5. FORMULÁRIO PRINCIPAL — CISPR 15
    # ══════════════════════════════════════════════════════════════════════
    section_bar(doc, '5. Formulário Principal — CISPR 15')
    body_para(doc,
        'O formulário principal é o ponto de entrada para emissão de relatórios CISPR 15. '
        'Acessa-se pela rota /cispr15 do sistema. É composto por quatro abas: '
        'Formulário, Clientes, Emendas e Relatórios. O ícone de engrenagem no canto '
        'superior direito dá acesso direto às Configurações.')

    heading_para(doc, '5.1 Aba Formulário', size=11, bold=True, space_before=8, space_after=3)
    body_para(doc, 'A aba Formulário é dividida nas seguintes seções:')

    bullet(doc, 'Tipo de DUT: seleção entre Lâmpada ou Luminária. Para lâmpadas, permite configurar as tensões de ensaio (127 V, 127+220 V ou 127+220+277 V). O identificador do produto é Código de Barras para lâmpadas e Número de Série para luminárias.')
    bullet(doc, 'Cliente: campos de nome, endereço, cidade e CEP. Ao digitar o CEP completo, o campo de cidade é preenchido automaticamente via API ViaCEP. Permite buscar clientes já cadastrados no banco de dados e salvar novos clientes.')
    bullet(doc, 'Objeto Ensaiado: campos de produto, fabricante, modelo, identificador, potência, tensão de alimentação e frequência. Possui botão "Ler Fotos (OCR)" para extração automática de dados das etiquetas do equipamento.')
    bullet(doc, 'Dados do Relatório: número do relatório (preenchido automaticamente ao gerar), responsável técnico, orçamento LABELO, protocolo, período de ensaio, data de emissão e resultado dos ensaios (Conduzida, Loop, Anexo B — PASS/FAIL).')
    bullet(doc, 'Anexos: seleção da pasta do ensaio (contendo .docx do Radimation + subpasta de fotos). Exibe miniaturas das fotos carregadas com opções de trocar ou remover individualmente.')

    add_screenshot(doc, '01_formulario_principal.png',
        'Formulário CISPR 15 — parte superior: tipo de DUT, cliente e objeto ensaiado')
    add_screenshot(doc, '01b_formulario_inferior.png',
        'Formulário CISPR 15 — parte inferior: dados do relatório, resultados e botões de ação')

    heading_para(doc, '5.2 Lógica de Validação e Bloqueio', size=11, bold=True, space_before=8, space_after=3)
    body_para(doc,
        'O sistema realiza validação em tempo real de todos os campos obrigatórios antes de '
        'permitir a geração do relatório. Os campos obrigatórios são: nome do cliente, endereço, '
        'cidade, produto, fabricante, modelo, identificador, potência nominal, tensão de '
        'alimentação, protocolo LABELO, responsável técnico, fotos do ensaio e arquivo .docx. '
        'O botão "Gerar Relatório" é bloqueado enquanto houver campos pendentes; um banner '
        'exibe a lista de pendências.')
    body_para(doc,
        'Ao gerar o relatório, o sistema verifica duplicidade de protocolo tanto no histórico '
        'local quanto na planilha Excel. Se houver duplicata, o operador é alertado e pode '
        'confirmar ou cancelar o registro. Após a geração, o formulário é automaticamente '
        'bloqueado (banner laranja "Formulário bloqueado") para evitar edições acidentais. '
        'O desbloqueio pode ser feito manualmente ou com senha, dependendo da configuração.')

    heading_para(doc, '5.3 Leitura OCR das Fotos', size=11, bold=True, space_before=8, space_after=3)
    body_para(doc,
        'O botão "Ler Fotos (OCR)" analisa as fotos da amostra (posições 3 e 4 para lâmpadas; '
        'posição 3 para luminárias) e extrai automaticamente campos como: produto, fabricante, '
        'modelo, código de barras/série, potência, tensão e frequência. O sistema utiliza '
        'primariamente a API Windows.Media.Ocr (nativa do Windows) via IPC do Electron; '
        'em caso de falha, usa Tesseract.js como alternativa. As sugestões são exibidas em '
        'cards clicáveis para aplicação individual ou em bloco.')

    heading_para(doc, '5.4 Pré-carregamento da Agenda', size=11, bold=True, space_before=8, space_after=3)
    body_para(doc,
        'Ao sair do campo Protocolo LABELO, o sistema busca automaticamente na Agenda de '
        'Execução um item com o mesmo protocolo e pré-carrega os dados do cliente, produto, '
        'fabricante, modelo, tensão, potência e demais campos disponíveis, eliminando a '
        'redigitação de informações já cadastradas.')

    page_break(doc)

    # ══════════════════════════════════════════════════════════════════════
    # 6. ABA CLIENTES
    # ══════════════════════════════════════════════════════════════════════
    section_bar(doc, '6. Aba Clientes — Banco de Dados de Clientes')
    body_para(doc,
        'A aba Clientes permite gerenciar o banco de dados de clientes do laboratório. '
        'Os dados são sincronizados com a pasta de rede configurada '
        '(arquivo cispr15_clientes.json), tornando-os disponíveis em todos os computadores '
        'do laboratório.')
    body_para(doc, 'Funcionalidades da aba Clientes:')
    bullet(doc, 'Listar todos os clientes cadastrados com nome, endereço, cidade e CEP.')
    bullet(doc, 'Buscar clientes por nome ou cidade.')
    bullet(doc, 'Utilizar um cliente: preenche automaticamente os campos do formulário principal.')
    bullet(doc, 'Editar dados de um cliente existente (nome, endereço, cidade, CEP, CNPJ).')
    bullet(doc, 'Excluir um cliente do banco de dados.')
    bullet(doc, 'Salvar no banco: o botão "Salvar no banco" no formulário principal salva o cliente atual ou atualiza se já existir.')

    add_screenshot(doc, '02_aba_clientes.png',
        'Aba Clientes — listagem e gerenciamento do banco de dados de clientes')

    page_break(doc)

    # ══════════════════════════════════════════════════════════════════════
    # 7. ABA EMENDAS
    # ══════════════════════════════════════════════════════════════════════
    section_bar(doc, '7. Aba Emendas — Histórico de Alterações de Relatórios')
    body_para(doc,
        'A aba Emendas exibe o histórico de emendas emitidas sobre relatórios já finalizados. '
        'Uma emenda é uma correção ou atualização formal de um relatório de ensaio já emitido, '
        'gerando um novo número de documento no formato "EMC XXXX/AAAA · Emenda N".')
    body_para(doc, 'Funcionalidades da aba Emendas:')
    bullet(doc, 'Listagem de todos os relatórios que possuem emendas emitidas, com o número do relatório original, data e alterações registradas.')
    bullet(doc, 'Carregar relatório: carrega os dados do relatório original para o formulário para geração da emenda.')
    bullet(doc, 'Excluir emenda: remove a emenda do histórico e apaga a cópia do PDF correspondente da pasta de cópias.')
    bullet(doc, 'Filtro por resultado: permite filtrar emendas por resultado (pass/fail) dos ensaios.')

    add_screenshot(doc, '03_aba_emendas.png',
        'Aba Emendas — histórico de alterações e emendas emitidas')

    body_para(doc,
        'Para gerar uma emenda, o operador acessa o botão "Gerar Emenda" no formulário '
        'principal. Se o formulário estiver bloqueado, o sistema solicita a senha de emissão '
        'antes de prosseguir. A tela de emenda permite registrar as alterações realizadas, '
        'que são listadas na última página do relatório PDF.')

    page_break(doc)

    # ══════════════════════════════════════════════════════════════════════
    # 8. ABA RELATÓRIOS
    # ══════════════════════════════════════════════════════════════════════
    section_bar(doc, '8. Aba Relatórios — Histórico de Relatórios Emitidos')
    body_para(doc,
        'A aba Relatórios exibe o histórico de todos os relatórios emitidos, carregados tanto '
        'do armazenamento local quanto da pasta de rede. Permite localizar e recarregar qualquer '
        'relatório anterior.')
    body_para(doc, 'Funcionalidades da aba Relatórios:')
    bullet(doc, 'Listagem com número do relatório, data de emissão, cliente, protocolo e produto.')
    bullet(doc, 'Busca por qualquer campo do relatório.')
    bullet(doc, 'Carregar relatório: recarrega todos os dados (configurações, fotos, .docx) para o formulário para visualização ou emissão de emenda.')
    bullet(doc, 'Ver PDF: abre diretamente a visualização do relatório em formato PDF sem recarregar o formulário.')
    bullet(doc, 'Exclusão de relatórios do histórico local.')

    add_screenshot(doc, '04_aba_relatorios.png',
        'Aba Relatórios — histórico de relatórios emitidos')

    page_break(doc)

    # ══════════════════════════════════════════════════════════════════════
    # 9. VISUALIZAÇÃO E GERAÇÃO DO RELATÓRIO PDF
    # ══════════════════════════════════════════════════════════════════════
    section_bar(doc, '9. Visualização e Geração do Relatório PDF')
    body_para(doc,
        'A tela de visualização do relatório (/cispr15/relatorio) exibe o relatório '
        'completo no formato A4, replicando fielmente o layout do documento impresso. '
        'O PDF é gerado diretamente pelo processo Electron (printToPDF), garantindo '
        'fidelidade total ao que é exibido na tela, sem dependência de impressora.')

    body_para(doc, 'Estrutura do relatório gerado:')
    bullet(doc, 'Página 1 — Capa: cabeçalho institucional PUCRS/LABELO, identificação do relatório, período de ensaio, dados do cliente, objeto ensaiado, documentação e normas utilizadas.')
    bullet(doc, 'Páginas 2–4 — Limites normativos: tabelas de limites conduzidos (9 kHz–30 MHz) e radiados (9 kHz–30 MHz e 30 MHz–300 MHz) conforme NBR IEC/CISPR 15/2014.')
    bullet(doc, 'Páginas seguintes — Resultados Radimation: conteúdo do arquivo .docx do Radimation, convertido e inserido com cabeçalho em cada página.')
    bullet(doc, 'Página de incertezas: tabela de incertezas de medição com fator de abrangência k=2.')
    bullet(doc, 'Páginas de fotos: duas fotos por página, com legendas (mínimo 4 slots/2 páginas).')
    bullet(doc, 'Página de emenda (quando aplicável): histórico de alterações com marcadores numerados em vermelho.')
    bullet(doc, 'Última página — Observações finais: 8 observações padrão do laboratório e linha para assinatura do signatário autorizado.')
    body_para(doc,
        'Ao clicar em "Baixar PDF" na barra de controles, o Electron usa printToPDF para '
        'gerar o arquivo com nome padronizado (NumRelatorio_tipo_fabricante.pdf) e salvá-lo '
        'na pasta da EUT. Uma cópia é automaticamente enviada para a pasta de cópias de PDF, '
        'se configurada.')

    add_screenshot(doc, '09_relatorio_pdf.png',
        'Tela de visualização do relatório — layout A4 com cabeçalho e rodapé')

    page_break(doc)

    # ══════════════════════════════════════════════════════════════════════
    # 10. EMISSÃO EM LOTE
    # ══════════════════════════════════════════════════════════════════════
    section_bar(doc, '10. Emissão em Lote')
    body_para(doc,
        'O sistema suporta emissão de múltiplos relatórios de um mesmo cliente em uma única '
        'operação (lote), acessível pelo botão "Emitir Lote" no formulário principal ou pelo '
        'botão "Gerar Lote" na Agenda de Execução.')
    body_para(doc, 'Funcionalidades da emissão em lote:')
    bullet(doc, 'Definição do número de amostras, tipo (lâmpada/luminária) e dados comuns (cliente, responsável).')
    bullet(doc, 'Dados individuais por amostra: produto, fabricante, modelo, identificador, potência, tensão, protocolo e orçamento.')
    bullet(doc, 'Carregamento de pasta de ensaio e fotos por amostra.')
    bullet(doc, 'Geração sequencial de todos os relatórios do lote com registro automático na planilha Excel.')
    bullet(doc, 'Integração com a Agenda: protocolos da agenda podem ser selecionados em lote para geração agrupada.')

    page_break(doc)

    # ══════════════════════════════════════════════════════════════════════
    # 11. AGENDA DE EXECUÇÃO
    # ══════════════════════════════════════════════════════════════════════
    section_bar(doc, '11. Agenda de Execução')
    body_para(doc,
        'A Agenda de Execução é a ferramenta de gerenciamento e acompanhamento do fluxo de '
        'ensaios do laboratório. Acessível pela rota /agenda, é independente do formulário '
        'de emissão e sempre acessível sem senha.')

    heading_para(doc, '11.1 Aba Agenda', size=11, bold=True, space_before=8, space_after=3)
    body_para(doc, 'Exibe todos os itens cadastrados com os seguintes dados:')
    bullet(doc, 'Tipo de equipamento (lâmpada/luminária), protocolo, orçamento, cliente, produto, datas de entrada e previsão de saída.')
    bullet(doc, 'Indicador visual de prazo: verde (em dia), amarelo (vence em ≤ 3 dias), vermelho (vencido).')
    bullet(doc, 'Badges de status dos ensaios: C (Conduzida), L (Loop), B (Anexo B) — clicáveis para alternar entre Pendente e Realizado.')
    bullet(doc, 'Para itens concluídos: exibe o número do relatório como link e botão para abrir o PDF.')
    bullet(doc, 'Filtros por status (em andamento/concluídos/todos), busca por texto e filtro por cliente.')
    bullet(doc, 'Ordenação por data de entrada, previsão de saída ou protocolo.')
    bullet(doc, 'Marcadores (tags): Urgente, Reensaio, Reprov. Conduzida, Reprov. Loop, Reprov. Anexo B.')

    add_screenshot(doc, '06_agenda.png',
        'Agenda de Execução — aba Agenda com listagem de itens e indicadores de prazo')
    add_screenshot(doc, '07_agenda_novo_item.png',
        'Agenda — modal de cadastro de novo item com dados completos do protocolo')

    heading_para(doc, '11.2 Aba Busca', size=11, bold=True, space_before=8, space_after=3)
    body_para(doc,
        'Permite busca unificada em toda a agenda e no histórico de relatórios emitidos, '
        'pesquisando por protocolo, orçamento, cliente, produto ou número do relatório.')

    heading_para(doc, '11.3 Aba Análise', size=11, bold=True, space_before=8, space_after=3)
    body_para(doc,
        'Painel analítico com indicadores de desempenho do laboratório:')
    bullet(doc, 'Cards de resumo: total, em andamento, concluídos, urgentes (≤ 3 dias) e vencidos.')
    bullet(doc, 'Distribuição de status: gráfico de barras horizontais por categoria (em dia, urgente, vencido).')
    bullet(doc, 'Entradas por mês: gráfico dos últimos 6 meses de entrada de equipamentos.')
    bullet(doc, 'Top clientes: clientes com maior número de ensaios.')
    bullet(doc, 'Lista de itens vencidos e urgentes com datas destacadas.')
    bullet(doc, 'Contagem de marcadores (tags) em uso.')

    add_screenshot(doc, '07_agenda_analise.png',
        'Agenda — aba Análise com indicadores de desempenho e distribuição por prazo')

    heading_para(doc, '11.4 Integração Agenda ↔ Formulário', size=11, bold=True, space_before=8, space_after=3)
    body_para(doc,
        'Ao gerar um relatório no formulário principal, o sistema busca automaticamente na '
        'agenda o item com o mesmo protocolo e registra o número do relatório e a data de '
        'emissão, atualizando seu status para "concluído" sem necessidade de ação manual. '
        'O botão "Ir para protocolo" na agenda pré-carrega os dados do item no formulário.')

    page_break(doc)

    # ══════════════════════════════════════════════════════════════════════
    # 12. FLUXO DE TRABALHO PADRÃO
    # ══════════════════════════════════════════════════════════════════════
    section_bar(doc, '12. Fluxo de Trabalho Padrão')
    body_para(doc,
        'O fluxo recomendado para emissão de um relatório CISPR 15 utilizando o sistema é '
        'o seguinte:')

    passos = [
        ('1', 'Cadastro na Agenda',          'O recebimento da amostra é registrado na Agenda de Execução com protocolo, orçamento, cliente e previsão de saída.'),
        ('2', 'Realização dos Ensaios',      'O técnico realiza os ensaios no Radimation e exporta o arquivo .docx. As fotos da amostra são organizadas na pasta da EUT (subpasta com numeração).'),
        ('3', 'Abertura do Formulário',      'Ao abrir o sistema, o formulário é limpo automaticamente. O operador preenche o protocolo LABELO — os dados da agenda são pré-carregados automaticamente.'),
        ('4', 'Carregamento da Pasta',       'O operador clica em "Selecionar Pasta do Ensaio" e seleciona a pasta da EUT. O sistema carrega o .docx e as fotos automaticamente.'),
        ('5', 'Revisão dos Dados',           'O operador revisa e complementa os dados do formulário. Opcionalmente utiliza "Ler Fotos (OCR)" para extrair dados da etiqueta do produto.'),
        ('6', 'Definição dos Resultados',    'O operador define o resultado dos ensaios (PASS/FAIL) para cada item: Conduzida, Loop e Anexo B.'),
        ('7', 'Geração do Relatório',        'O operador clica em "Gerar Relatório". O sistema verifica duplicatas, registra na planilha Excel, bloqueia o formulário e exibe a tela de visualização PDF.'),
        ('8', 'Download/Salvar PDF',         'O operador clica em "Baixar PDF". O PDF é salvo na pasta da EUT e copiado para a pasta de cópias, se configurada.'),
        ('9', 'Atualização da Agenda',       'O sistema atualiza automaticamente o item da agenda com o número do relatório e a data de emissão.'),
    ]

    tbl = doc.add_table(rows=len(passos), cols=3)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, (num, etapa, desc) in enumerate(passos):
        c0, c1, c2 = tbl.cell(i, 0), tbl.cell(i, 1), tbl.cell(i, 2)
        c0.width = Cm(1.0)
        c1.width = Cm(4.5)
        c2.width = Cm(12.5)
        set_cell_bg(c0, BG_DARK)
        p0 = c0.paragraphs[0]; p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p0, num, bold=True, size=11, color=WHITE)
        set_cell_bg(c1, BG_GRAY1)
        p1 = c1.paragraphs[0]; p1.paragraph_format.left_indent = Cm(0.2)
        add_run(p1, etapa, bold=True, size=9.5)
        row_bg = BG_ALT if i % 2 == 0 else 'FFFFFF'
        set_cell_bg(c2, row_bg)
        p2 = c2.paragraphs[0]; p2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY; p2.paragraph_format.left_indent = Cm(0.2)
        add_run(p2, desc, size=9.5)

    doc.add_paragraph()
    page_break(doc)

    # ══════════════════════════════════════════════════════════════════════
    # 13. CRITÉRIOS DE VALIDAÇÃO
    # ══════════════════════════════════════════════════════════════════════
    section_bar(doc, '13. Critérios de Validação do Software')
    body_para(doc,
        'A validação do software foi realizada conforme os requisitos da ABNT NBR ISO/IEC 17025:2017, '
        'item 6.4.13 (validação de software para uso em laboratório). A tabela a seguir '
        'apresenta os critérios, testes realizados e resultados obtidos:')

    headers_val = ['Item', 'Critério de Validação', 'Teste Realizado', 'Resultado Esperado', 'Status']
    criterios = [
        ('VAL-01', 'Integridade dos dados do formulário',
         'Preencher todos os campos obrigatórios e verificar persistência após recarregamento da página.',
         'Dados preservados no localStorage entre recarregamentos.',
         'APROVADO'),
        ('VAL-02', 'Validação de campos obrigatórios',
         'Tentar gerar relatório com campos pendentes.',
         'Sistema bloqueia geração e lista campos pendentes.',
         'APROVADO'),
        ('VAL-03', 'Registro automático na planilha Excel',
         'Gerar relatório e verificar registro na planilha de controle.',
         'Número sequencial gerado e registrado na planilha com dados corretos.',
         'APROVADO'),
        ('VAL-04', 'Geração do PDF com layout correto',
         'Gerar PDF de relatório de lâmpada e de luminária e verificar estrutura.',
         'PDF com todas as páginas (capa, limites, Radimation, fotos, observações) corretamente formatadas.',
         'APROVADO'),
        ('VAL-05', 'Carregamento de pasta da EUT',
         'Selecionar pasta contendo .docx e subpasta de fotos.',
         'Sistema carrega .docx e todas as imagens, preenchendo campos de protocolo automaticamente.',
         'APROVADO'),
        ('VAL-06', 'OCR de etiquetas',
         'Acionar OCR com foto nítida de etiqueta de produto.',
         'Campos de fabricante, modelo, tensão e potência preenchidos com dados extraídos.',
         'APROVADO'),
        ('VAL-07', 'Sincronização via pasta de rede',
         'Configurar pasta de rede e verificar acesso a dados de dois computadores distintos.',
         'Clientes e relatórios visíveis em ambos os computadores.',
         'APROVADO'),
        ('VAL-08', 'Controle de acesso por senha',
         'Configurar senha e tentar acessar formulário sem autenticação.',
         'Modal de senha exibido; acesso negado com senha incorreta; acesso liberado com senha correta.',
         'APROVADO'),
        ('VAL-09', 'Emissão de emenda',
         'Gerar emenda sobre relatório existente e verificar geração do PDF com marcadores.',
         'PDF de emenda com número correto (EMC XXXX/AAAA · Emenda N) e tabela de alterações.',
         'APROVADO'),
        ('VAL-10', 'Bloqueio de formulário após emissão',
         'Gerar relatório e tentar editar campos.',
         'Formulário bloqueado (banner laranja) após geração; campos não editáveis.',
         'APROVADO'),
        ('VAL-11', 'Verificação de protocolo duplicado',
         'Tentar registrar protocolo já existente na planilha.',
         'Sistema exibe aviso de duplicata com número do relatório anterior; operador confirma ou cancela.',
         'APROVADO'),
        ('VAL-12', 'Pré-carregamento da Agenda',
         'Cadastrar item na agenda e preencher protocolo no formulário.',
         'Dados do item (cliente, produto, fabricante, etc.) carregados automaticamente.',
         'APROVADO'),
        ('VAL-13', 'Emissão em lote',
         'Gerar lote com 3 amostras de mesmo cliente.',
         'Três relatórios gerados com numeração sequencial e registro correto na planilha.',
         'APROVADO'),
        ('VAL-14', 'Cópia automática de PDF',
         'Gerar relatório com pasta de cópias configurada.',
         'Cópia do PDF salva automaticamente na pasta configurada.',
         'APROVADO'),
        ('VAL-15', 'Integridade do banco de clientes',
         'Cadastrar, editar e excluir clientes; verificar persistência na pasta de rede.',
         'Dados de clientes atualizados corretamente no JSON compartilhado.',
         'APROVADO'),
    ]

    tbl_v = doc.add_table(rows=len(criterios) + 1, cols=5)
    tbl_v.style = 'Table Grid'
    tbl_v.alignment = WD_TABLE_ALIGNMENT.LEFT
    # Header
    widths = [Cm(1.6), Cm(4.5), Cm(4.5), Cm(4.0), Cm(2.2)]
    for j, h in enumerate(headers_val):
        c = tbl_v.cell(0, j)
        c.width = widths[j]
        set_cell_bg(c, BG_DARK)
        p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p, h, bold=True, size=8.5, color=WHITE)
    # Linhas
    for i, (item, crit, teste, esperado, status) in enumerate(criterios, 1):
        vals = [item, crit, teste, esperado, status]
        row_bg = BG_ALT if i % 2 == 0 else 'FFFFFF'
        for j, val in enumerate(vals):
            c = tbl_v.cell(i, j)
            set_cell_bg(c, row_bg)
            p = c.paragraphs[0]
            p.paragraph_format.left_indent = Cm(0.15)
            bold = (j == 0)
            color = DARK if j == 4 and val == 'APROVADO' else BLACK
            add_run(p, val, size=8, bold=bold, color=color)

    doc.add_paragraph()
    page_break(doc)

    # ══════════════════════════════════════════════════════════════════════
    # 14. ARMAZENAMENTO E ARQUITETURA DE DADOS
    # ══════════════════════════════════════════════════════════════════════
    section_bar(doc, '14. Arquitetura de Armazenamento de Dados')
    body_para(doc,
        'O sistema utiliza uma arquitetura híbrida de armazenamento, combinando '
        'armazenamento local (localStorage do Electron / navegador) com arquivos JSON '
        'em pasta de rede compartilhada:')

    arq = [
        ('settings.json',               '%AppData%\\CISPR15-LABELO\\',  'Configurações do sistema (caminhos, senha). Local por computador.'),
        ('cispr15_clientes.json',        'Pasta de rede configurada',   'Banco de dados de clientes. Compartilhado entre todos os PCs.'),
        ('cispr15_relatorios.json',      'Pasta de rede configurada',   'Metadados de relatórios (sem fotos e sem HTML do docx). Compartilhado.'),
        ('cispr15_agenda.json',          'Pasta da agenda configurada', 'Itens da agenda de execução. Compartilhado.'),
        ('localStorage',                 'Computador local',            'Formulário atual (cfg), fotos (base64), HTML do docx, histórico local completo.'),
        ('[numRelatorio]_tipo_fab.pdf',  'Pasta da EUT + Pasta cópias', 'PDF do relatório. Salvo na pasta do ensaio e copiado para pasta de cópias.'),
    ]

    tbl_a = doc.add_table(rows=len(arq) + 1, cols=3)
    tbl_a.style = 'Table Grid'
    for j, h in enumerate(['Arquivo / Local', 'Localização', 'Conteúdo']):
        c = tbl_a.cell(0, j)
        set_cell_bg(c, BG_DARK)
        p = c.paragraphs[0]; p.paragraph_format.left_indent = Cm(0.2)
        add_run(p, h, bold=True, size=9.5, color=WHITE)
    for i, (arqv, loc, cont) in enumerate(arq, 1):
        row_bg = BG_ALT if i % 2 == 0 else 'FFFFFF'
        for j, val in enumerate([arqv, loc, cont]):
            c = tbl_a.cell(i, j)
            set_cell_bg(c, row_bg)
            p = c.paragraphs[0]; p.paragraph_format.left_indent = Cm(0.2)
            add_run(p, val, size=9, bold=(j == 0))

    doc.add_paragraph()

    # ══════════════════════════════════════════════════════════════════════
    # 15. MÓDULO LAB — VISÃO GERAL
    # ══════════════════════════════════════════════════════════════════════
    page_break(doc)
    section_bar(doc, '15. Módulo Lab — Gestão de Equipamentos e Certificados')
    body_para(doc,
        'O Módulo Lab é um conjunto de ferramentas complementares ao sistema CISPR 15 LABELO, '
        'voltado para a gestão do laboratório EMC: cadastro de equipamentos de medição, '
        'certificados de calibração com grades de correção 2D, checagens intermediárias de '
        'instrumentos e elaboração de instruções de trabalho / procedimentos de calibração. '
        'O módulo é acessado pela barra lateral esquerda do sistema e inclui as rotas '
        '/dashboard, /equipamentos, /certificados, /checagens e /procedimentos.')

    heading_para(doc, '15.1 Dashboard', size=11, bold=True, space_before=8, space_after=3)
    body_para(doc,
        'A tela de Dashboard (/dashboard) exibe uma visão consolidada do laboratório '
        'com quatro indicadores: total de equipamentos cadastrados, normas ativas, '
        'checagens pendentes (vencendo em ≤ 30 dias) e checagens vencidas. '
        'Também lista as próximas cinco checagens em ordem de prazo com seus status.')

    add_screenshot(doc, 'lab_dashboard.png',
        'Dashboard do Módulo Lab — indicadores gerais e próximas checagens')

    # ══════════════════════════════════════════════════════════════════════
    # 16. EQUIPAMENTOS
    # ══════════════════════════════════════════════════════════════════════
    page_break(doc)
    section_bar(doc, '16. Equipamentos — Cadastro e Gestão')
    body_para(doc,
        'A tela de Equipamentos (/equipamentos) permite cadastrar e gerenciar todos os '
        'instrumentos do laboratório EMC. Cada equipamento possui uma TAG identificadora '
        'única e é classificado por grupo e subgrupo.')

    heading_para(doc, '16.1 Grupos e Subgrupos', size=11, bold=True, space_before=8, space_after=3)
    body_para(doc,
        'Os equipamentos são organizados nos seguintes grupos:')
    bullet(doc, 'Geradores de sinal — ex.: geradores de RF, fontes de ruído.')
    bullet(doc, 'Medidores — receptores EMI, multímetros, osciloscópios.')
    bullet(doc, 'Redes de Impedância (LISN) — redes de impedância de linha.')
    bullet(doc, 'Antenas — antenas de ensaio e calibradas.')
    bullet(doc, 'Atenuação — atenuadores, cabos, acopladores.')
    bullet(doc, 'Grandezas Ambientais — termômetros, higrômetros.')

    heading_para(doc, '16.2 Dados do Equipamento', size=11, bold=True, space_before=8, space_after=3)
    body_para(doc, 'Cada equipamento registra:')
    bullet(doc, 'TAG: identificador único (ex.: GEN-001, ANT-003).')
    bullet(doc, 'Nome, fabricante, modelo e número de série.')
    bullet(doc, 'Grupo e subgrupo de classificação.')
    bullet(doc, 'Data de próxima calibração e status (Ativo / Calibrar / Fora de uso).')
    bullet(doc, 'Certificados de calibração vinculados.')

    add_screenshot(doc, 'lab_equipamentos.png',
        'Tela de Equipamentos — lista por grupo com indicador de status')

    # ══════════════════════════════════════════════════════════════════════
    # 17. CERTIFICADOS DE CALIBRAÇÃO
    # ══════════════════════════════════════════════════════════════════════
    page_break(doc)
    section_bar(doc, '17. Certificados de Calibração — Grade 2D')
    body_para(doc,
        'A tela de Certificados (/certificados) gerencia os certificados de calibração '
        'acreditados (ABNT NBR ISO/IEC 17025) dos padrões do laboratório. Cada certificado '
        'pode conter uma grade de correção 2D usada para interpolação bilinear nas checagens '
        'intermediárias.')

    heading_para(doc, '17.1 Importação via PDF (OCR)', size=11, bold=True, space_before=8, space_after=3)
    body_para(doc,
        'O sistema aceita certificados em PDF (ex.: LABELO/PUCRS, modelo Agilent N5171B). '
        'O arquivo é processado com pdf-parse com agrupamento espacial de pixels, '
        'extraindo as tabelas de calibração com separação de colunas por tabulação. '
        'A página de capa é ignorada automaticamente; todas as demais páginas são processadas, '
        'incluindo tabelas de distorção harmônica e de varredura de frequência.')
    body_para(doc,
        'O parser identifica automaticamente:')
    bullet(doc, 'Número do certificado e TAG do equipamento.')
    bullet(doc, 'Cabeçalhos de tabela com VR (valor de referência / UST) e MM (valor medido / UMP).')
    bullet(doc, 'Parâmetros (ex.: "Linearidade", "Distorção Harmônica") como nome da tabela.')
    bullet(doc, '"Medição de ..." como grandeza (ex.: "Potência", "Frequência").')
    bullet(doc, 'Correção calculada automaticamente: Correção = VR − MM.')
    bullet(doc, 'Incerteza (IM) detectada no cabeçalho ou pela heurística de coluna adjacente ao MM.')
    body_para(doc,
        'A sub-linha de unidades ("UST(dBm) / UMP(dBm)") é ignorada pelo parser, '
        'evitando a sobrescrição do cabeçalho correto. Números no formato brasileiro '
        '("1.000,5" = 1000,5 ; "2.450" = 2450) são reconhecidos automaticamente.',
        italic=True)

    heading_para(doc, '17.2 Grade de Correção 2D', size=11, bold=True, space_before=8, space_after=3)
    body_para(doc,
        'A grade 2D representa uma função de correção f(Frequência, Nível) ou f(VR, Nível), '
        'com pontos agrupados por parâmetro (tabela). Para uso nas checagens, o sistema '
        'realiza interpolação bilinear sobre a grade para calcular a correção no ponto '
        'exato de frequência e nível solicitado, mesmo que não conste explicitamente no '
        'certificado. A tela exibe os pontos em tabela com colunas: Frequência (eixo 1), '
        'Nível VR (eixo 2), MM, Correção (VR−MM) e Incerteza.')

    add_screenshot(doc, 'lab_certificado_grade.png',
        'Certificado com grade 2D — tabela de pontos por parâmetro e grade de correção')

    # ══════════════════════════════════════════════════════════════════════
    # 18. CHECAGENS INTERMEDIÁRIAS
    # ══════════════════════════════════════════════════════════════════════
    page_break(doc)
    section_bar(doc, '18. Checagens Intermediárias — Controle de Instrumentos')
    body_para(doc,
        'A tela de Checagens (/checagens) controla as verificações intermediárias periódicas '
        'dos instrumentos de medição, conforme exigido pelo sistema da qualidade '
        'ABNT NBR ISO/IEC 17025:2017. As checagens são organizadas em três categorias: '
        'Vencidas, Vencendo nos próximos 30 dias e Em dia.')

    heading_para(doc, '18.1 Tipos de Checagem', size=11, bold=True, space_before=8, space_after=3)
    body_para(doc, 'Existem dois modos de registro:')
    bullet(doc, 'Do certificado: seleciona um ponto já existente no certificado do padrão. O valor de referência (VR) e a incerteza são preenchidos automaticamente a partir da grade 2D.')
    bullet(doc, 'Manual com interpolação: o operador define um ponto não calibrado (frequência e nível). O sistema interpola bilinearmente a correção e calcula o VR automaticamente.')

    heading_para(doc, '18.2 Dados de uma Checagem', size=11, bold=True, space_before=8, space_after=3)
    body_para(doc, 'Cada checagem registra:')
    bullet(doc, 'TAG do equipamento verificado e certificado do padrão utilizado.')
    bullet(doc, 'Data de realização e data da próxima checagem.')
    bullet(doc, 'Norma de referência (ex.: ABNT NBR ISO/IEC 17025).')
    bullet(doc, 'Valor de referência (VR), valor medido (MM) e correção aplicada.')
    bullet(doc, 'Status calculado automaticamente: APROVADO, ATENÇÃO ou REPROVADO.')
    bullet(doc, 'Observações do técnico.')

    heading_para(doc, '18.3 Templates de Checagem', size=11, bold=True, space_before=8, space_after=3)
    body_para(doc,
        'A tela de Templates (/checagens/templates) permite criar modelos reutilizáveis de '
        'checagem, pré-configurados com equipamento, parâmetros e periodicidade. '
        'Facilita o registro de checagens recorrentes com um clique.')

    add_screenshot(doc, 'lab_checagens.png',
        'Tela de Checagens — lista por status com indicador de prazo')

    # ══════════════════════════════════════════════════════════════════════
    # 19. NORMAS
    # ══════════════════════════════════════════════════════════════════════
    page_break(doc)
    section_bar(doc, '19. Normas — Banco de Normas Técnicas')
    body_para(doc,
        'A tela de Normas (/normas) gerencia o banco de normas técnicas utilizadas nos '
        'ensaios e calibrações do laboratório. Permite cadastrar normas com número, '
        'título, organismo publicador (ABNT, IEC, CISPR, etc.), edição, ano e '
        'situação (vigente, substituída, cancelada). '
        'O dashboard exibe o total de normas ativas como indicador de referência.')

    # ══════════════════════════════════════════════════════════════════════
    # 20. PROCEDIMENTOS — IT e PC
    # ══════════════════════════════════════════════════════════════════════
    page_break(doc)
    section_bar(doc, '20. Procedimentos — Instruções de Trabalho e Procedimentos de Calibração')
    body_para(doc,
        'O módulo de Procedimentos (/procedimentos) centraliza a criação e gestão de '
        'Instruções de Trabalho (IT) e Procedimentos de Calibração (PC) no formato '
        'padrão LABELO.')

    heading_para(doc, '20.1 Editor de IT / PC', size=11, bold=True, space_before=8, space_after=3)
    body_para(doc,
        'O editor de documentos (/procedimentos/instrucoes/novo) oferece um ambiente '
        'estruturado para elaboração de documentos técnicos com os seguintes elementos:')
    bullet(doc, 'Seções numeradas com hierarquia (1, 1.1, 1.1.1).')
    bullet(doc, 'Parágrafos de texto justificado.')
    bullet(doc, 'Listas com marcadores e listas de passos numerados.')
    bullet(doc, 'Imagens embutidas com legenda.')
    bullet(doc, 'Tabelas com cabeçalho e linhas alternadas.')
    bullet(doc, 'Blocos de destaque (NOTA, ATENÇÃO, AVISO).')
    bullet(doc, 'Siglas e definições com glossário automático.')

    heading_para(doc, '20.2 Lista de Documentos', size=11, bold=True, space_before=8, space_after=3)
    body_para(doc,
        'A tela /procedimentos/instrucoes lista todos os documentos cadastrados com número, '
        'título, revisão e data. Documentos podem ser editados, revisados (incremento '
        'automático do número de revisão) ou excluídos.')

    add_screenshot(doc, 'lab_procedimentos.png',
        'Módulo Procedimentos — visão geral com acesso a checagens e editor de IT/PC')

    # ══════════════════════════════════════════════════════════════════════
    # 21. SIDEBAR DE NAVEGAÇÃO DO MÓDULO LAB
    # ══════════════════════════════════════════════════════════════════════
    page_break(doc)
    section_bar(doc, '21. Navegação — Sidebar do Módulo Lab')
    body_para(doc,
        'A barra lateral (sidebar) do Módulo Lab provê acesso rápido a todos os recursos. '
        'É exibida em todas as telas do módulo e organizada em seções:')

    tbl_nav = doc.add_table(rows=9, cols=2)
    tbl_nav.style = 'Table Grid'
    tbl_nav.alignment = WD_TABLE_ALIGNMENT.LEFT
    nav_itens = [
        ('Dashboard',                  '/dashboard — visão consolidada com indicadores e próximas checagens'),
        ('Equipamentos',               '/equipamentos — lista e cadastro de instrumentos por grupo'),
        ('Certificados',               '/certificados — certificados de calibração com grade 2D'),
        ('Checagens',                  '/checagens — registro e controle de checagens intermediárias'),
        ('Normas',                     '/normas — banco de normas técnicas do laboratório'),
        ('Procedimentos',              '/procedimentos — hub de checagens e editor de IT/PC'),
        ('Instruções de Trabalho',     '/procedimentos/instrucoes — lista de documentos IT/PC'),
        ('Relatórios CISPR 15',        '/cispr15 — retorna ao formulário principal de emissão'),
    ]
    for j, h in enumerate(['Seção', 'Descrição']):
        c = tbl_nav.cell(0, j)
        set_cell_bg(c, BG_DARK)
        p = c.paragraphs[0]; p.paragraph_format.left_indent = Cm(0.2)
        add_run(p, h, bold=True, size=9.5, color=WHITE)
    for i, (sec, desc) in enumerate(nav_itens, 1):
        row_bg = BG_ALT if i % 2 == 0 else 'FFFFFF'
        set_cell_bg(tbl_nav.cell(i, 0), row_bg)
        set_cell_bg(tbl_nav.cell(i, 1), row_bg)
        p0 = tbl_nav.cell(i, 0).paragraphs[0]; p0.paragraph_format.left_indent = Cm(0.2)
        add_run(p0, sec, bold=True, size=9)
        p1 = tbl_nav.cell(i, 1).paragraphs[0]; p1.paragraph_format.left_indent = Cm(0.2)
        add_run(p1, desc, size=9)

    doc.add_paragraph()
    page_break(doc)

    # ══════════════════════════════════════════════════════════════════════
    # 22. FUNÇÕES RECENTES (ATUALIZAÇÕES)
    # ══════════════════════════════════════════════════════════════════════
    page_break(doc)
    section_bar(doc, '22. Funções Recentes (Atualizações)')
    body_para(doc,
        'Esta seção documenta as funcionalidades acrescentadas nas últimas versões do '
        'sistema, com foco na gestão metrológica (Módulo Lab) e na agenda. As demais '
        'funções descritas nos capítulos anteriores permanecem inalteradas.')

    heading_para(doc, '22.1 Importação em lote de certificados (pasta-mãe)', size=11, bold=True, space_before=8, space_after=3)
    body_para(doc,
        'Em Equipamentos → "Importar pasta-mãe", o sistema percorre uma pasta com uma '
        'subpasta por equipamento (cada uma com o certificado em PDF), lê os dados e '
        'cadastra automaticamente. Os certificados do LABELO são cadastrados direto; os de '
        'outros laboratórios vão para o Rascunho, identificados pelo laboratório emissor.')
    bullet(doc, 'Equipamentos parados há mais de 7 anos sem alteração na pasta (provável fora de uso): '
                'o sistema pergunta se deseja cadastrá-los, enviá-los ao Rascunho ou ignorá-los.')
    bullet(doc, '2ª varredura ("Cadastrar pela amostra"): cadastra os itens do Rascunho pelos dados '
                'da folha de rosto, mesmo sem o certificado padrão do LABELO.')

    heading_para(doc, '22.2 Laboratórios de Calibração e Modelo de Extração', size=11, bold=True, space_before=8, space_after=3)
    body_para(doc,
        'A tela Laboratórios mantém o vínculo entre a acreditação (CAL XXXX do selo azul '
        'ABNT NBR ISO/IEC 17025) e o nome do laboratório. O LABELO é identificado '
        'exclusivamente pelo seu CAL 0024 — a mera presença da palavra "LABELO" (que aparece '
        'como cliente em certificados de terceiros) não classifica o documento como do LABELO.')
    bullet(doc, 'Importar certificados: lê PDFs e associa automaticamente o CAL ao nome do laboratório emissor.')
    bullet(doc, 'Modelo de extração por laboratório: para cada lab, é possível informar qual RÓTULO '
                'ele usa em cada campo (Nome, Fabricante, Modelo, Série, TAG, Data). O botão '
                '"Importar amostra" exibe o texto extraído do PDF para localizar os rótulos. '
                'O OCR passa a usar esses rótulos com prioridade para aquele laboratório.')

    heading_para(doc, '22.3 Siglas oficiais e identificação da TAG', size=11, bold=True, space_before=8, space_after=3)
    body_para(doc,
        'A TAG do equipamento segue o padrão número + 3 letras (ex.: 1528EMC). As 3 letras '
        'devem ser uma sigla oficial de laboratório cadastrada (Taxonomia → Siglas); qualquer '
        'outra trinca de letras não é considerada TAG, evitando leituras incorretas. As siglas '
        'oficiais já vêm cadastradas e podem ser vinculadas a cada laboratório/área.')

    heading_para(doc, '22.4 Cadastro pela Análise Crítica (FOR 6401)', size=11, bold=True, space_before=8, space_after=3)
    body_para(doc,
        'Quando o PDF lido é um formulário FOR 6401 (Análise Crítica de Certificado de '
        'Calibração), o sistema extrai diretamente dele: Fornecedor (laboratório), número do '
        'Certificado, TAG, Nome do Instrumento, Data do certificado e Periodicidade. '
        'Periodicidades em anos são convertidas para meses. Quando há mais de uma análise '
        'crítica para a mesma TAG, prevalece a periodicidade da análise mais recente.')

    heading_para(doc, '22.5 Grandezas dos certificados do LABELO', size=11, bold=True, space_before=8, space_after=3)
    body_para(doc,
        'Ao cadastrar um certificado do LABELO, as grandezas são identificadas automaticamente '
        'a partir dos títulos de seção do certificado (textos centralizados/negrito que '
        'antecedem cada "Parâmetro:", entre a 2ª e a penúltima página) e registradas no '
        'equipamento, ficando disponíveis no seletor de grandeza das checagens.')

    heading_para(doc, '22.6 Tipos de equipamento (atribuição de grupos)', size=11, bold=True, space_before=8, space_after=3)
    body_para(doc,
        'Em Equipamentos → Grupos, o painel "Tipos de equipamento" lista todos os nomes '
        'distintos de equipamentos cadastrados e permite atribuir cada nome a um grupo e '
        'subgrupo de uma só vez (vale para todos os equipamentos com aquele nome), agilizando '
        'a classificação correta.')

    heading_para(doc, '22.7 Agenda — fluxo de assinatura e custódia', size=11, bold=True, space_before=8, space_after=3)
    bullet(doc, 'Estados do item: "Em andamento" → ao emitir vai para "Aguardando assinatura" → '
                'ao marcar como assinado vai para "Concluído" (com a data registrada).')
    bullet(doc, 'O botão do relatório abre a PASTA do relatório (DOCX, fotos e PDF) para permitir '
                'a assinatura manual.')
    bullet(doc, 'Cadeia de custódia da amostra: caixas de confirmação de recebimento na entrega '
                '(EMC) e na devolução (LUM).')

    heading_para(doc, '22.8 Pasta de cópias de PDF e arquivos do relatório', size=11, bold=True, space_before=8, space_after=3)
    bullet(doc, 'A cópia do PDF para a pasta de cópias ocorre em apenas duas situações: ao Assinar '
                'e Publicar pelo sistema, ou quando o PDF original (na pasta do DOCX) é assinado '
                'manualmente — detectado ao reabrir/visualizar o relatório. Apenas gerar o PDF não cria cópia.')
    bullet(doc, '"Salvar arquivos" vincula as fotos e o DOCX ao relatório (sem duplicar arquivos na '
                'pasta); ao reabrir o protocolo, eles voltam para dentro do PDF.')
    bullet(doc, '"Baixar PDF" salva na pasta do DOCX (pasta da EUT), e não em Documentos.')

    heading_para(doc, '22.9 Correções e desempenho', size=11, bold=True, space_before=8, space_after=3)
    bullet(doc, 'Status "Fora de uso", "Não requer calibração" e "Calibrar antes do uso" não exibem data de próxima calibração.')
    bullet(doc, 'O botão "Excluir tudo" (Equipamentos) passou a funcionar com um modal de senha próprio.')
    bullet(doc, 'Desempenho: a gravação de dados deixou de travar a digitação — a interface não congela mais ao salvar.')

    # ══════════════════════════════════════════════════════════════════════
    # 23. OBSERVAÇÕES FINAIS E ASSINATURAS
    # ══════════════════════════════════════════════════════════════════════
    page_break(doc)
    section_bar(doc, '23. Declaração de Validação e Assinaturas')
    body_para(doc,
        'Com base nos testes realizados e documentados neste manual, declara-se que o '
        'software CISPR 15 LABELO atende aos requisitos operacionais do laboratório e está '
        'em conformidade com os critérios de validação de software estabelecidos pela '
        'ABNT NBR ISO/IEC 17025:2017, sendo adequado para uso na geração de relatórios de '
        'ensaio de compatibilidade eletromagnética conforme a norma NBR IEC/CISPR 15/2014.')

    body_para(doc, f'Porto Alegre, {TODAY}.', justify=False, before=10, after=10)

    # Tabela de assinaturas
    tbl_s = doc.add_table(rows=2, cols=3)
    tbl_s.alignment = WD_TABLE_ALIGNMENT.CENTER
    colunas = ['Elaborado por', 'Revisado por', 'Aprovado por']
    nomes   = ['Dionata Rafael Blauth\nda Paixão Nunes', 'Jonathan Culau', 'Laboratório\nLABELO-PUCRS']
    for j in range(3):
        c_top = tbl_s.cell(0, j)
        c_bot = tbl_s.cell(1, j)
        set_cell_bg(c_top, BG_DARK)
        p_top = c_top.paragraphs[0]; p_top.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p_top, colunas[j], bold=True, size=9, color=WHITE)
        p_bot = c_bot.paragraphs[0]; p_bot.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_bot.paragraph_format.space_before = Pt(20)
        p_bot.paragraph_format.space_after  = Pt(4)
        add_run(p_bot, nomes[j], size=9)

    doc.add_paragraph()

# ─── configuração de página ───────────────────────────────────────────────
def setup_page(doc):
    from docx.oxml import OxmlElement
    section = doc.sections[0]
    section.page_width  = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.0)
    section.right_margin  = Cm(1.5)

# ─── main ─────────────────────────────────────────────────────────────────
def main():
    doc = Document()
    setup_page(doc)
    build_header(doc)
    build_footer(doc)
    build_capa(doc)
    build_body(doc)
    doc.save(str(OUT))
    print(f'\nDocumento salvo em:\n   {OUT}')

if __name__ == '__main__':
    main()
