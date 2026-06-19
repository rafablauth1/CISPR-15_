# -*- coding: utf-8 -*-
"""Substitui os placeholders "[Imagem não encontrada: X]" no manual pelas imagens
capturadas em screenshots_manual/."""
import re, os
import docx
from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

BASE  = r"C:\Users\Notla\OneDrive\Área de Trabalho\DIONATA\cispr15-standalone"
DOC   = os.path.join(BASE, "Manual_Validacao_CISPR15_LABELO.docx")
SHOTS = os.path.join(BASE, "screenshots_manual")

d = docx.Document(DOC)
trocadas = 0
for p in d.paragraphs:
    m = re.search(r'encontrada:\s*([^\]]+)\]', p.text)
    if not m:
        continue
    fname = m.group(1).strip()
    fp = os.path.join(SHOTS, fname)
    if not os.path.exists(fp):
        print('faltando ainda:', fname); continue
    for r in list(p.runs):
        r.text = ''
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(fp, width=Cm(15.5))
    trocadas += 1
    print('injetada:', fname)

d.save(DOC)
print(f'\n{trocadas} imagem(ns) injetada(s) em {DOC}')
